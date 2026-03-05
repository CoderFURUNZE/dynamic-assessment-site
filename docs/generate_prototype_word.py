from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import patches
from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
FIG_DIR = ROOT / "figures_prototype"
OUT_DOC = ROOT / "system_prototype_rich.docx"


def setup_style():
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def _new_ax(title: str):
    fig, ax = plt.subplots(figsize=(12, 7), dpi=160)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.02, 0.96, title, fontsize=18, fontweight="bold", va="top")
    return fig, ax


def _box(ax, x, y, w, h, text, fc="#E8F0FE", ec="#4A6CF7", fontsize=11):
    rect = patches.FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.01,rounding_size=0.02",
        linewidth=1.6,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(rect)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize)


def _arrow(ax, x1, y1, x2, y2, text: str | None = None):
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(arrowstyle="->", lw=1.7, color="#2B3A67"),
    )
    if text:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.02, text, fontsize=10, ha="center")


def fig_01_context():
    fig, ax = _new_ax("图1  系统上下文图")
    _box(ax, 0.05, 0.6, 0.18, 0.14, "学生")
    _box(ax, 0.05, 0.35, 0.18, 0.14, "管理员/教师")
    _box(ax, 0.36, 0.46, 0.28, 0.2, "动态评价与推荐系统")
    _box(ax, 0.78, 0.63, 0.16, 0.12, "PostgreSQL")
    _box(ax, 0.78, 0.42, 0.16, 0.12, "视频资源")
    _box(ax, 0.78, 0.21, 0.16, 0.12, "行为信号流")
    _arrow(ax, 0.23, 0.67, 0.36, 0.57, "学习请求")
    _arrow(ax, 0.23, 0.42, 0.36, 0.55, "配置管理")
    _arrow(ax, 0.64, 0.60, 0.78, 0.69, "读写")
    _arrow(ax, 0.64, 0.52, 0.78, 0.48, "拉取")
    _arrow(ax, 0.78, 0.27, 0.64, 0.50, "反馈")
    return fig


def fig_02_architecture():
    fig, ax = _new_ax("图2  系统分层架构图")
    _box(ax, 0.08, 0.77, 0.84, 0.12, "表示层：Vue3 + Element Plus（学生端/管理端）", fc="#FFF3E0", ec="#FB8C00")
    _box(ax, 0.08, 0.59, 0.84, 0.12, "接口层：FastAPI REST + WebSocket", fc="#E3F2FD", ec="#1E88E5")
    _box(ax, 0.08, 0.37, 0.84, 0.16, "业务层：认证/课程/知识点/题库/小测/练习/推荐/报表/审计", fc="#E8F5E9", ec="#43A047")
    _box(ax, 0.08, 0.15, 0.84, 0.16, "数据层：PostgreSQL + 媒体目录 + 日志", fc="#F3E5F5", ec="#8E24AA")
    _arrow(ax, 0.50, 0.77, 0.50, 0.71)
    _arrow(ax, 0.50, 0.59, 0.50, 0.53)
    _arrow(ax, 0.50, 0.37, 0.50, 0.31)
    return fig


def fig_03_module_tree():
    fig, ax = _new_ax("图3  功能模块结构图")
    _box(ax, 0.38, 0.84, 0.24, 0.1, "系统功能")
    nodes = [
        ("账号权限", 0.08, 0.63),
        ("学习端", 0.30, 0.63),
        ("管理端", 0.52, 0.63),
        ("推荐引擎", 0.74, 0.63),
    ]
    for t, x, y in nodes:
        _box(ax, x, y, 0.18, 0.1, t, fc="#EEF4FF")
        _arrow(ax, 0.50, 0.84, x + 0.09, 0.73)
    _box(ax, 0.27, 0.45, 0.24, 0.08, "资源/小测/练习/笔记", fc="#E9F7EF", ec="#43A047")
    _box(ax, 0.49, 0.45, 0.24, 0.08, "课程/知识点/题库/报表", fc="#E9F7EF", ec="#43A047")
    _box(ax, 0.71, 0.45, 0.24, 0.08, "规则策略+模型打分", fc="#E9F7EF", ec="#43A047")
    return fig


def fig_04_usecase_student():
    fig, ax = _new_ax("图4  学生端用例图（简化）")
    actor = patches.Circle((0.08, 0.55), 0.05, facecolor="#FFFDE7", edgecolor="#F9A825", lw=1.5)
    ax.add_patch(actor)
    ax.text(0.08, 0.55, "学生", ha="center", va="center", fontsize=11)
    _box(ax, 0.24, 0.73, 0.2, 0.08, "查看学习资源")
    _box(ax, 0.24, 0.61, 0.2, 0.08, "完成小测")
    _box(ax, 0.24, 0.49, 0.2, 0.08, "完成练习")
    _box(ax, 0.24, 0.37, 0.2, 0.08, "查看掌握度")
    _box(ax, 0.24, 0.25, 0.2, 0.08, "错题复习")
    for y in [0.77, 0.65, 0.53, 0.41, 0.29]:
        _arrow(ax, 0.13, 0.55, 0.24, y)
    return fig


def fig_05_usecase_admin():
    fig, ax = _new_ax("图5  管理端用例图（简化）")
    actor = patches.Circle((0.08, 0.55), 0.05, facecolor="#FFFDE7", edgecolor="#F9A825", lw=1.5)
    ax.add_patch(actor)
    ax.text(0.08, 0.55, "管理员", ha="center", va="center", fontsize=11)
    items = ["课程管理", "知识点管理", "先修关系管理", "题库与小测管理", "报表与日志查看"]
    for i, txt in enumerate(items):
        y = 0.77 - i * 0.12
        _box(ax, 0.24, y, 0.28, 0.08, txt, fc="#F3E5F5", ec="#8E24AA")
        _arrow(ax, 0.13, 0.55, 0.24, y + 0.04)
    return fig


def fig_06_er():
    fig, ax = _new_ax("图6  核心 E-R 关系图（简化）")
    _box(ax, 0.05, 0.72, 0.18, 0.1, "Course")
    _box(ax, 0.30, 0.72, 0.20, 0.1, "KnowledgePoint")
    _box(ax, 0.58, 0.78, 0.18, 0.1, "Question")
    _box(ax, 0.58, 0.62, 0.18, 0.1, "Quiz/QuizItem")
    _box(ax, 0.82, 0.70, 0.14, 0.1, "Resource")
    _box(ax, 0.30, 0.46, 0.20, 0.1, "User")
    _box(ax, 0.58, 0.46, 0.18, 0.1, "PracticeAttempt")
    _box(ax, 0.82, 0.46, 0.14, 0.1, "ExpressionEvent")
    _box(ax, 0.58, 0.30, 0.18, 0.1, "Mastery")
    _arrow(ax, 0.23, 0.77, 0.30, 0.77, "1..N")
    _arrow(ax, 0.50, 0.77, 0.58, 0.83, "1..N")
    _arrow(ax, 0.50, 0.77, 0.58, 0.67, "1..N")
    _arrow(ax, 0.50, 0.77, 0.82, 0.75, "1..N")
    _arrow(ax, 0.50, 0.51, 0.58, 0.51, "1..N")
    _arrow(ax, 0.76, 0.51, 0.82, 0.51, "1..N")
    _arrow(ax, 0.67, 0.46, 0.67, 0.40, "1..1")
    return fig


def fig_07_reco_flow():
    fig, ax = _new_ax("图7  推荐流程图")
    _box(ax, 0.06, 0.78, 0.22, 0.1, "输入：做题记录/掌握度")
    _box(ax, 0.35, 0.78, 0.22, 0.1, "规则评估\n证据清单+阈值")
    _box(ax, 0.64, 0.78, 0.22, 0.1, "模型评分\nMLP概率")
    _box(ax, 0.20, 0.55, 0.22, 0.1, "候选题生成")
    _box(ax, 0.50, 0.55, 0.22, 0.1, "融合排序")
    _box(ax, 0.35, 0.32, 0.28, 0.1, "输出：下一题/解锁建议")
    _arrow(ax, 0.28, 0.83, 0.35, 0.83)
    _arrow(ax, 0.57, 0.83, 0.64, 0.83)
    _arrow(ax, 0.46, 0.78, 0.31, 0.65)
    _arrow(ax, 0.75, 0.78, 0.61, 0.65)
    _arrow(ax, 0.42, 0.60, 0.50, 0.60)
    _arrow(ax, 0.61, 0.55, 0.49, 0.42)
    return fig


def fig_08_sequence():
    fig, ax = _new_ax("图8  练习提交流程序列图（简化）")
    x_student, x_front, x_api, x_db = 0.12, 0.34, 0.56, 0.80
    for x, name in [(x_student, "学生"), (x_front, "前端"), (x_api, "后端API"), (x_db, "数据库")]:
        ax.text(x, 0.88, name, ha="center", fontsize=11, fontweight="bold")
        ax.plot([x, x], [0.18, 0.84], color="#90A4AE", lw=1.2, linestyle="--")
    y = 0.80
    def msg(x1, x2, text):
        nonlocal y
        _arrow(ax, x1, y, x2, y, text)
        y -= 0.08
    msg(x_student, x_front, "提交答案")
    msg(x_front, x_api, "POST /practice/submit")
    msg(x_api, x_db, "写入 PracticeAttempt")
    msg(x_db, x_api, "返回结果")
    msg(x_api, x_db, "更新 Mastery")
    msg(x_api, x_front, "返回 正确性+解析")
    msg(x_front, x_student, "展示反馈")
    return fig


def fig_09_deploy():
    fig, ax = _new_ax("图9  部署拓扑图（开发环境）")
    _box(ax, 0.07, 0.60, 0.22, 0.12, "浏览器\nhttp://localhost:5173", fc="#FFF8E1", ec="#FFA000")
    _box(ax, 0.39, 0.72, 0.22, 0.12, "Vue Dev Server\nVite :5173", fc="#E3F2FD", ec="#1E88E5")
    _box(ax, 0.39, 0.50, 0.22, 0.12, "FastAPI\nUvicorn :8000", fc="#E8F5E9", ec="#43A047")
    _box(ax, 0.72, 0.62, 0.22, 0.12, "PostgreSQL\nDocker :5432", fc="#F3E5F5", ec="#8E24AA")
    _box(ax, 0.72, 0.40, 0.22, 0.12, "pgAdmin\nDocker :5050", fc="#FCE4EC", ec="#D81B60")
    _arrow(ax, 0.29, 0.66, 0.39, 0.78, "前端访问")
    _arrow(ax, 0.50, 0.72, 0.50, 0.62, "API请求")
    _arrow(ax, 0.61, 0.56, 0.72, 0.68, "SQL")
    _arrow(ax, 0.83, 0.40, 0.83, 0.62, "管理")
    return fig


def fig_10_roadmap():
    fig, ax = _new_ax("图10  阶段迭代路线图")
    ax.plot([0.1, 0.9], [0.5, 0.5], color="#78909C", lw=2)
    phases = [
        (0.15, "V1\n基础学习闭环"),
        (0.35, "V2\n推荐策略完善"),
        (0.55, "V3\n报表与审计"),
        (0.75, "V4\n行为信号融合"),
        (0.90, "V5\n部署与优化"),
    ]
    for x, t in phases:
        ax.add_patch(patches.Circle((x, 0.5), 0.02, color="#4A6CF7"))
        ax.text(x, 0.42, t, ha="center", va="top", fontsize=10)
    return fig


def save_fig(fig, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    p = FIG_DIR / name
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def build_figures() -> list[Path]:
    setup_style()
    figs = [
        save_fig(fig_01_context(), "fig01_context.png"),
        save_fig(fig_02_architecture(), "fig02_architecture.png"),
        save_fig(fig_03_module_tree(), "fig03_module_tree.png"),
        save_fig(fig_04_usecase_student(), "fig04_usecase_student.png"),
        save_fig(fig_05_usecase_admin(), "fig05_usecase_admin.png"),
        save_fig(fig_06_er(), "fig06_er.png"),
        save_fig(fig_07_reco_flow(), "fig07_reco_flow.png"),
        save_fig(fig_08_sequence(), "fig08_sequence.png"),
        save_fig(fig_09_deploy(), "fig09_deploy.png"),
        save_fig(fig_10_roadmap(), "fig10_roadmap.png"),
    ]
    return figs


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = 1
    p.runs[0].font.size = Pt(10.5)


def build_doc(figs: list[Path]):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)

    doc.add_heading("动态评价与自适应推荐学习系统", level=0)
    doc.add_paragraph("课程设计/毕业设计系统原型（图文版）")
    doc.add_paragraph("说明：本文档基于当前 Learning 项目实现情况自动生成，侧重系统原型与结构表达。")

    doc.add_heading("1. 系统概述", level=1)
    doc.add_paragraph(
        "系统面向课程学习过程中的个性化推荐场景，提供学习端与管理端。学习端支持资源学习、小测、练习、错题复习；"
        "管理端支持课程、知识点、先修关系、题库、小测、报表和审计管理。"
    )

    doc.add_heading("2. 总体架构与模块", level=1)
    for idx in [0, 1, 2]:
        doc.add_picture(str(figs[idx]), width=Inches(6.7))
        add_caption(doc, f"图{idx+1}  系统结构示意")

    doc.add_heading("3. 角色与用例", level=1)
    for idx in [3, 4]:
        doc.add_picture(str(figs[idx]), width=Inches(6.7))
        add_caption(doc, f"图{idx+1}  角色用例示意")

    doc.add_heading("4. 数据模型", level=1)
    doc.add_picture(str(figs[5]), width=Inches(6.9))
    add_caption(doc, "图6  核心实体关系示意")
    doc.add_paragraph("核心实体包括 User、Course、KnowledgePoint、Question、PracticeAttempt、Mastery、ExpressionEvent 等。")

    doc.add_heading("5. 推荐系统原型", level=1)
    doc.add_picture(str(figs[6]), width=Inches(6.9))
    add_caption(doc, "图7  推荐流程示意")
    doc.add_paragraph(
        "推荐引擎采用规则策略与模型评分融合：先依据掌握度阈值、证据清单和难度步进生成候选，再结合模型概率进行排序。"
    )

    doc.add_heading("6. 核心业务流程", level=1)
    doc.add_picture(str(figs[7]), width=Inches(6.9))
    add_caption(doc, "图8  练习提交流程序列示意")

    doc.add_heading("7. 部署原型", level=1)
    doc.add_picture(str(figs[8]), width=Inches(6.9))
    add_caption(doc, "图9  本地开发部署拓扑")

    doc.add_heading("8. 迭代路线", level=1)
    doc.add_picture(str(figs[9]), width=Inches(6.9))
    add_caption(doc, "图10  版本迭代路线")

    doc.add_heading("9. 结论", level=1)
    doc.add_paragraph(
        "当前原型已覆盖“学习闭环 + 管理闭环 + 推荐闭环”核心能力，后续可继续提升推荐评估、界面交互一致性与部署稳定性。"
    )

    doc.save(OUT_DOC)


def main():
    figs = build_figures()
    build_doc(figs)
    print(f"Generated: {OUT_DOC}")
    print(f"Figures dir: {FIG_DIR}")


if __name__ == "__main__":
    main()

