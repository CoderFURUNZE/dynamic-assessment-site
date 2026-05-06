from pathlib import Path

from docx import Document


template = Path(r"C:\Users\34286\Desktop\软件学院毕业设计中期报告模板（软件学院2026届-v2）.docx")
source = Path(r"C:\Users\34286\Desktop\毕业设计相关资料\融合学习者画像与知识图谱的动态评价模型研究及应用-中期报告.docx")

for label, path in (("template", template), ("source", source)):
    print(label, path, path.exists())
    doc = Document(path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "nonempty", len(texts))
    for text in texts[:30]:
        print(text)
    for index, table in enumerate(doc.tables):
        print("TABLE", index, len(table.rows), len(table.columns))
        for row in table.rows[:10]:
            print(" | ".join(cell.text.replace("\n", " / ")[:160] for cell in row.cells))
    print("---")
